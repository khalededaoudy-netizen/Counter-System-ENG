"""Fills the ticket number into the existing Word template.

The template's fonts/colors/images/layout are left completely alone —
this only rewrites the text of the placeholder run(s) that spell
"{{TICKET_NUMBER}}", wherever it appears (body paragraphs, tables,
headers/footers).

Template requirement: type the placeholder in one continuous edit
(don't let AutoCorrect/AutoFormat split it into pieces) so it lands in
a single run and its original formatting (font/size/color) is kept
exactly. If Word did split it across runs anyway, this still works but
the merged span collapses to the formatting of its first run — call
out to the template author to retype the placeholder if that happens.
"""
from __future__ import annotations

import shutil
import uuid
from pathlib import Path

from docx import Document

PLACEHOLDER = "{{TICKET_NUMBER}}"


class TemplateError(Exception):
    pass


def _replace_in_paragraph(paragraph, replacement: str) -> bool:
    full_text = "".join(r.text for r in paragraph.runs)
    if PLACEHOLDER not in full_text:
        return False

    new_text = full_text.replace(PLACEHOLDER, replacement)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


def _replace_everywhere(doc: Document, replacement: str) -> int:
    count = 0

    def walk_paragraphs(paragraphs):
        nonlocal count
        for p in paragraphs:
            if _replace_in_paragraph(p, replacement):
                count += 1

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    walk_paragraphs(cell.paragraphs)
                    walk_tables(cell.tables)

    walk_paragraphs(doc.paragraphs)
    walk_tables(doc.tables)
    for section in doc.sections:
        for part in (section.header, section.footer):
            walk_paragraphs(part.paragraphs)
            walk_tables(part.tables)

    return count


def fill_ticket_number(
    template_path: str | Path,
    ticket_number: int,
    output_dir: str | Path,
    number_padding: int = 3,
) -> Path:
    template_path = Path(template_path)
    if not template_path.exists():
        raise TemplateError(f"Ticket template not found: {template_path}")

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    display_number = str(ticket_number).zfill(number_padding) if number_padding else str(ticket_number)

    doc = Document(str(template_path))
    replaced = _replace_everywhere(doc, display_number)
    if replaced == 0:
        raise TemplateError(
            f"Placeholder '{PLACEHOLDER}' was not found anywhere in the template."
        )

    out_path = output_dir / f"ticket_{display_number}_{uuid.uuid4().hex[:8]}.docx"
    doc.save(str(out_path))
    return out_path


def cleanup(path: str | Path) -> None:
    try:
        Path(path).unlink(missing_ok=True)
    except OSError:
        pass
