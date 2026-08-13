from docx import Document

from app.printing.template_filler import PLACEHOLDER, TemplateError, fill_ticket_number


def _make_template(path, text=None):
    doc = Document()
    doc.add_paragraph(text if text is not None else f"Ticket number: {PLACEHOLDER}")
    doc.save(str(path))


def test_fill_replaces_placeholder(tmp_path):
    template_path = tmp_path / "template.docx"
    _make_template(template_path)

    out = fill_ticket_number(template_path, 7, tmp_path / "out", number_padding=3)
    assert out.exists()

    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "007" in full_text
    assert PLACEHOLDER not in full_text


def test_fill_without_padding(tmp_path):
    template_path = tmp_path / "template.docx"
    _make_template(template_path)

    out = fill_ticket_number(template_path, 42, tmp_path / "out", number_padding=0)
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "42" in full_text


def test_missing_placeholder_raises(tmp_path):
    template_path = tmp_path / "template.docx"
    _make_template(template_path, text="No placeholder here")

    try:
        fill_ticket_number(template_path, 1, tmp_path / "out")
        assert False, "expected TemplateError"
    except TemplateError:
        pass


def test_missing_template_file_raises(tmp_path):
    try:
        fill_ticket_number(tmp_path / "nope.docx", 1, tmp_path / "out")
        assert False, "expected TemplateError"
    except TemplateError:
        pass
